import os
from openai import OpenAI
import pandas as pd
import json
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def generar_quiz(tema, grado, cantidad, tipo, formato):
    # Si es College Board, fuerza Word (aunque el front lo bloquee, doble seguridad)
    if grado.strip().lower() == "college board" or tipo.strip().lower() == "college board":
        return generar_word_collegeboard(tema, grado, cantidad)

    if formato == "kahoot":
        return generar_excel(tema, grado, cantidad)
    else:
        return generar_word(tema, grado, cantidad, tipo)



def generar_excel(tema, grado, cantidad):
    prompt = f"""
Crea {cantidad} preguntas de opción múltiple para estudiantes de {grado} basate en los estandares del departamento de educacion en Puerto Rico.
Tema: {tema}

Devuelve SOLO JSON:
[
  {{
    "pregunta": "",
    "A": "",
    "B": "",
    "C": "",
    "D": "",
    "correcta": "A"
  }}
]
"""

    response = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt
    )

    texto = response.output_text.strip()
    texto = texto.replace("```json", "").replace("```", "")
    preguntas = json.loads(texto)

    letra = {"A": 1, "B": 2, "C": 3, "D": 4}

    df = pd.DataFrame({
        "Question": [p["pregunta"] for p in preguntas],
        "Answer 1": [p["A"] for p in preguntas],
        "Answer 2": [p["B"] for p in preguntas],
        "Answer 3": [p["C"] for p in preguntas],
        "Answer 4": [p["D"] for p in preguntas],
        "Time Limit": [0 for _ in preguntas],
        "Correct answer": [letra[p["correcta"]] for p in preguntas]
    })

    archivo = "quiz_kahoot.xlsx"
    df.to_excel(archivo, index=False)
    return archivo


def comic_sans(run, size=12):
    run.font.name = "Comic Sans MS"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Comic Sans MS")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Comic Sans MS")
    run._element.rPr.rFonts.set(qn("w:cs"), "Comic Sans MS")
    run.font.size = Pt(size)

def generar_word(tema, grado, cantidad, tipo):
    prompt = f"""
Crea una prueba escrita para estudiantes de {grado} basate en los estandares del departamento de educacion en Puerto Rico(no incluyas en la prueba que son de ahi).
Tema: {tema}
Cantidad de preguntas: {cantidad}
Tipo de preguntas: {tipo}
Formato limpio y académico (sin emojis, sin símbolos raros)
No incluyas fecha, nombre del maestro, ni referencias externas

REGLAS IMPORTANTES:
1) Si el tema es de Inglés (English / Grammar / Reading / Writing / ...), escribe TODA la evaluación en INGLÉS.
2) Si el tema es de Español (Español / Gramatica / Lectura / Escritura / ...), escribe TODA la evaluación en ESPAÑOL.
3) Si el hay DOS o mas temas (ejemplo: "Fracciones y Pronouns" o "Fracciones, Pronouns, ..."):
   - Divide la evaluación en 2 secciones por idioma.
   - Sección 1 en ESPAÑOL (temas en español).
   - Sección 2 en INGLÉS (temas en inglés).
4) No mezcles idiomas dentro de una misma sección.

Incluye:
- Título
- Instrucciones, incluye en ellas que debe poner sus respuestas y procedimiento en papel aparte.(solo si aplica)
- Preguntas numeradas
"""

    response = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt
    )

    texto = response.output_text.strip()

    doc = Document()

    for linea in texto.split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(linea)
        comic_sans(r, 12)

    archivo = "prueba.docx"
    doc.save(archivo)
    return archivo

def generar_word_collegeboard(tema, grado, cantidad):
    prompt = f"""
Crea una evaluación estilo College Board.
Tema(s): {tema}
Cantidad total de preguntas: {cantidad}

REGLAS IMPORTANTES:
1) Si el tema es de Inglés (English / Grammar / Reading / Writing / ...), escribe TODA la evaluación en INGLÉS.
2) Si el tema es de Español (Español / Gramatica / Lectura / Escritura / ...), escribe TODA la evaluación en ESPAÑOL.
3) Si el hay DOS o mas temas (ejemplo: "Fracciones y Pronouns" o "Fracciones, Pronouns, ..."):
   - Divide la evaluación en 2 secciones por idioma.
   - Sección 1 en ESPAÑOL (temas en español).
   - Sección 2 en INGLÉS (temas en inglés).
4) No mezcles idiomas dentro de una misma sección.

FORMATO ESTILO COLLEGE BOARD:
- Formato limpio y académico (sin emojis, sin símbolos raros)
- Incluye secciones (por ejemplo: Multiple Choice, Short Response, Free Response) según aplique al tema
- En opción múltiple: 4 opciones (A-D)
- En respuesta corta o ensayo: indica claramente qué debe incluir la respuesta
- No incluyas fecha, nombre del maestro, ni referencias externas

Devuelve SOLO el contenido de la prueba en texto, con preguntas numeradas y secciones claras.
"""

    response = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt
    )

    texto = response.output_text.strip()
    doc = Document()

    for linea in texto.split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(linea)
        comic_sans(r, 12)

    archivo = "prueba_collegeboard.docx"
    doc.save(archivo)
    return archivo

